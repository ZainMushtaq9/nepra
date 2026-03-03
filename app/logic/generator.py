# app/logic/generator.py
from docx import Document
from docx.shared import Pt, Inches
import json
import os
from datetime import datetime
from flask import current_app

def generate_complaint_docx(bill):
    """
    Generates a DOCX file containing a formal complaint to NEPRA/Wafaqi Mohtasib.
    """
    doc = Document()
    
    # Try parsing bill JSON safely
    try:
        data = json.loads(bill.bill_json)
    except:
        data = {}

    # Extract Data with Defaults
    consumer_name = data.get("Consumer Name") or "[Your Name]"
    ref_no = data.get("Reference No") or "[Reference Number]"
    issue_date = data.get("Issue Date") or "[Bill Issue Date]"
    total_amount = data.get("Total Amount") or "[Amount]"
    
    # Document Styling
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Header
    doc.add_heading('FORMAL COMPLAINT REGARDING ELECTRICITY BILL DISCREPANCY', 1)
    
    # Date
    p_date = doc.add_paragraph(f"Date: {datetime.now().strftime('%d %B, %Y')}")
    p_date.alignment = 2 # Right align

    # To Address
    if bill.fault_type == "Company Fault":
        to_address = "To,\nThe Wafaqi Mohtasib (Ombudsman) Secretariat,\nRegional Office."
    else:
        to_address = "To,\nThe Sub-Divisional Officer (SDO),\nElectricity Distribution Company."
        
    doc.add_paragraph(to_address)
    
    # Subject
    p_sub = doc.add_paragraph()
    p_sub.add_run(f"Subject: COMPLAINT AGAINST OVERBILLING / DISCREPANCY FOR REFERENCE NO. {ref_no}").bold = True
    
    # Salutation
    doc.add_paragraph("\nRespected Sir/Madam,")
    
    # Body Paragraph 1 (Introduction)
    p_body = doc.add_paragraph(
        f"I, {consumer_name}, hold the electricity connection under Reference No. {ref_no}. "
        f"I am writing this formal complaint against the abnormal electricity bill issued to me for the period of {issue_date} "
        f"amounting to Rs. {total_amount}."
    )
    
    # Body Paragraph 2 (Analysis Result)
    doc.add_paragraph("\nAccording to an analysis based on the NEPRA Consumer Service Manual:")
    p_analysis = doc.add_paragraph(bill.analysis_result)
    p_analysis.paragraph_format.left_indent = Inches(0.5)
    
    # Body Paragraph 3 (Closing)
    doc.add_paragraph(
        "\nIn light of the above-mentioned points and specific NEPRA guidelines, I request your kind office to immediately "
        "intervene, rectify the bill, and issue a revised current bill without imposing late payment surcharges while the "
        "investigation is ongoing, as per my rights defined in the Consumer Service Manual."
    )
    
    doc.add_paragraph("\nLooking forward to your prompt response and justice.")
    
    # Footer
    doc.add_paragraph("\nYours sincerely,")
    doc.add_paragraph(f"\n{consumer_name}")
    doc.add_paragraph("Signature: ______________________")
    doc.add_paragraph(f"Connection Reference: {ref_no}")
    doc.add_paragraph("Contact No: _____________________")
    doc.add_paragraph("Address: _________________________________________________")
    
    # Save to file
    filename = f"Complaint_{ref_no}_{bill.id}.docx".replace(" ", "_")
    filepath = os.path.join(current_app.config['UPLOAD_FOLDER'], filename)
    doc.save(filepath)
    
    return filename, filepath
